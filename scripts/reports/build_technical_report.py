from __future__ import annotations

import csv
import json
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PAPER_DIR = ROOT / "docs" / "report"
ASSET_DIR = PAPER_DIR / "assets"
OUTPUT_DOCX = PAPER_DIR / "uav_vision_technical_report.docx"

FONT_REGULAR = Path("C:/Windows/Fonts/times.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/timesbd.ttf")
FONT_ITALIC = Path("C:/Windows/Fonts/timesi.ttf")

INK = "#111827"
BLUE = "#2563EB"
GREEN = "#059669"
ORANGE = "#D97706"
RED = "#DC2626"
GRAY = "#6B7280"
GRID = "#D1D5DB"
PANEL = "#F8FAFC"


def load_json(path: str) -> dict:
    return json.loads((ROOT / path).read_text(encoding="utf-8"))


def load_csv(path: str) -> list[dict[str, str]]:
    with (ROOT / path).open("r", encoding="utf-8", newline="") as file:
        return list(csv.DictReader(file))


def pil_font(size: int, bold: bool = False, italic: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD if bold else FONT_ITALIC if italic else FONT_REGULAR
    return ImageFont.truetype(str(path), size)


def draw_text_centered(
    draw: ImageDraw.ImageDraw,
    xy: tuple[float, float],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str = INK,
) -> None:
    box = draw.textbbox((0, 0), text, font=font)
    draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1]), text, font=font, fill=fill)


def draw_bar_panel(
    draw: ImageDraw.ImageDraw,
    bounds: tuple[int, int, int, int],
    title: str,
    labels: list[str],
    values: list[float],
    colors: list[str],
    value_format: str,
    maximum: float | None = None,
) -> None:
    x0, y0, x1, y1 = bounds
    draw.rounded_rectangle(bounds, radius=18, fill=PANEL, outline=GRID, width=2)
    draw_text_centered(draw, ((x0 + x1) / 2, y0 + 22), title, pil_font(38, bold=True))
    plot_left, plot_right = x0 + 105, x1 - 40
    plot_top, plot_bottom = y0 + 100, y1 - 110
    maximum = maximum or max(values) * 1.15
    for step in range(5):
        y = plot_bottom - (plot_bottom - plot_top) * step / 4
        value = maximum * step / 4
        draw.line((plot_left, y, plot_right, y), fill=GRID, width=2)
        draw.text((x0 + 18, y - 17), f"{value:.2f}", font=pil_font(24), fill=GRAY)
    slot = (plot_right - plot_left) / len(values)
    width = slot * 0.58
    for index, (label, value, color) in enumerate(zip(labels, values, colors)):
        cx = plot_left + slot * (index + 0.5)
        top = plot_bottom - (plot_bottom - plot_top) * value / maximum
        draw.rounded_rectangle((cx - width / 2, top, cx + width / 2, plot_bottom), radius=8, fill=color)
        draw_text_centered(draw, (cx, max(y0 + 72, top - 40)), value_format.format(value), pil_font(25, bold=True))
        label_lines = label.split("\n")
        for line_index, line in enumerate(label_lines):
            draw_text_centered(draw, (cx, plot_bottom + 16 + line_index * 28), line, pil_font(24))


def build_pipeline_figure() -> Path:
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    draw_text_centered(draw, (800, 35), "End-to-End UAV Vision Pipeline", pil_font(50, bold=True))

    labels = [
        ("Image / Video", BLUE),
        ("Resize and\nNormalize", GRAY),
        ("Slim YOLOv8s\nDetector", GREEN),
        ("ByteTrack\nAssociation", ORANGE),
        ("Boxes, IDs,\nand Video", BLUE),
    ]
    centers = [165, 475, 800, 1125, 1435]
    for index, ((label, color), center) in enumerate(zip(labels, centers)):
        draw.rounded_rectangle((center - 125, 190, center + 125, 370), radius=20, fill="#FFFFFF", outline=color, width=6)
        for line_index, line in enumerate(label.split("\n")):
            draw_text_centered(draw, (center, 235 + line_index * 48), line, pil_font(34, bold=line_index == 0))
        if index < len(labels) - 1:
            draw.line((center + 135, 280, centers[index + 1] - 145, 280), fill=INK, width=6)
            tip = centers[index + 1] - 145
            draw.polygon([(tip, 280), (tip - 25, 265), (tip - 25, 295)], fill=INK)

    draw.rounded_rectangle((255, 530, 1345, 790), radius=24, fill=PANEL, outline=GRID, width=3)
    draw_text_centered(draw, (800, 560), "Deployment Path", pil_font(38, bold=True))
    deployment = [("PyTorch", 430, BLUE), ("ONNX", 800, GRAY), ("TensorRT FP16", 1170, GREEN)]
    for index, (label, center, color) in enumerate(deployment):
        draw.rounded_rectangle((center - 120, 650, center + 120, 735), radius=18, fill="white", outline=color, width=5)
        draw_text_centered(draw, (center, 670), label, pil_font(32, bold=True))
        if index < len(deployment) - 1:
            next_center = deployment[index + 1][1]
            draw.line((center + 130, 692, next_center - 140, 692), fill=INK, width=5)
            draw.polygon([(next_center - 140, 692), (next_center - 165, 678), (next_center - 165, 706)], fill=INK)
    path = ASSET_DIR / "pipeline.png"
    image.save(path, quality=95)
    return path


def build_eda_figure(train: dict, val: dict) -> Path:
    image = Image.new("RGB", (1600, 920), "white")
    draw = ImageDraw.Draw(image)
    draw_text_centered(draw, (800, 28), "VisDrone Dataset Characteristics", pil_font(48, bold=True))
    categories = [(key, value) for key, value in train["category_distribution"].items() if key not in {"ignored", "others"}]
    categories = sorted(categories, key=lambda item: item[1], reverse=True)[:6]
    draw_bar_panel(
        draw,
        (30, 105, 1030, 880),
        "Top Training Categories (instances)",
        [item[0].replace("awning-tricycle", "awning\ntricycle") for item in categories],
        [float(item[1]) / 1000 for item in categories],
        [BLUE, GREEN, ORANGE, RED, "#7C3AED", "#0891B2"],
        "{:.1f}k",
        maximum=160,
    )
    draw_bar_panel(
        draw,
        (1060, 105, 1570, 880),
        "Objects < 32 x 32",
        ["Train", "Validation"],
        [
            float(train["small_object_ratio_lt_32x32"]),
            float(val["small_object_ratio_lt_32x32"]),
        ],
        [BLUE, ORANGE],
        "{:.1%}",
        maximum=0.8,
    )
    path = ASSET_DIR / "dataset_eda.png"
    image.save(path, quality=95)
    return path


def build_small_object_figure(rows: list[dict[str, str]]) -> Path:
    selected_keys = ["standard", "multiscale_768_960_1280", "sahi_640_overlap020", "focal_gamma2_e6_eval"]
    selected = [next(row for row in rows if row["key"] == key) for key in selected_keys]
    labels = ["Standard", "Multi-\nscale", "SAHI", "Focal"]
    colors = [GRAY, BLUE, GREEN, ORANGE]
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    draw_text_centered(draw, (800, 25), "Small-Object Accuracy-Speed Trade-off", pil_font(48, bold=True))
    draw_bar_panel(
        draw,
        (30, 100, 785, 870),
        "Small-Object Recall",
        labels,
        [float(row["small_recall"]) for row in selected],
        colors,
        "{:.3f}",
        maximum=0.75,
    )
    draw_bar_panel(
        draw,
        (815, 100, 1570, 870),
        "Throughput (FPS)",
        labels,
        [float(row["fps"]) for row in selected],
        colors,
        "{:.1f}",
        maximum=110,
    )
    path = ASSET_DIR / "small_object_tradeoff.png"
    image.save(path, quality=95)
    return path


def build_deployment_figure(calibration: list[dict[str, str]], runtime: list[dict[str, str]]) -> Path:
    calibration_by_name = {row["method"]: row for row in calibration}
    fp16 = calibration_by_name["FP16"]
    int8 = calibration_by_name["INT8 bright"]
    runtime_by_name = {row["name"]: row for row in runtime}
    labels = ["TensorRT\nFP16", "INT8\nbright", "FP16 +\nByteTrack"]
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    draw_text_centered(draw, (800, 25), "Deployment Accuracy and End-to-End Runtime", pil_font(48, bold=True))
    draw_bar_panel(
        draw,
        (30, 100, 785, 870),
        "Validation mAP50-95",
        ["FP16", "INT8\nbright"],
        [float(fp16["map50_95"]), float(int8["map50_95"])],
        [GREEN, ORANGE],
        "{:.3f}",
        maximum=0.32,
    )
    draw_bar_panel(
        draw,
        (815, 100, 1570, 870),
        "End-to-End Throughput (FPS)",
        labels,
        [
            float(runtime_by_name["slim_trt_fp16_960"]["fps"]),
            float(runtime_by_name["slim_trt_int8_960"]["fps"]),
            float(runtime_by_name["slim_trt_fp16_960_track"]["fps"]),
        ],
        [GREEN, ORANGE, BLUE],
        "{:.1f}",
        maximum=105,
    )
    path = ASSET_DIR / "deployment_tradeoff.png"
    image.save(path, quality=95)
    return path


def build_qualitative_figure() -> Path:
    detection = Image.open(ROOT / "docs/assets/yolo_distillation/distilled_val_predictions.jpg").convert("RGB")
    tracking = Image.open(
        ROOT / "outputs/tracking/yolov8s_slim04375_tensorrt_bytetrack_uav123_group1/visualizations/000300_track.jpg"
    ).convert("RGB")
    target_width = 1500
    detection.thumbnail((target_width, 760), Image.Resampling.LANCZOS)
    tracking.thumbnail((target_width, 760), Image.Resampling.LANCZOS)
    canvas = Image.new("RGB", (target_width + 60, detection.height + tracking.height + 170), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((30, 18), "(a) Distilled detector validation predictions", font=pil_font(38, bold=True), fill=INK)
    canvas.paste(detection, ((canvas.width - detection.width) // 2, 68))
    y = 90 + detection.height
    draw.text((30, y), "(b) TensorRT FP16 and ByteTrack output", font=pil_font(38, bold=True), fill=INK)
    canvas.paste(tracking, ((canvas.width - tracking.width) // 2, y + 50))
    path = ASSET_DIR / "qualitative_results.jpg"
    canvas.save(path, quality=92)
    return path


def set_run_font(run, size: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_columns(section, count: int, space_twips: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))


def set_cell_margins(cell, top: int = 35, start: int = 20, bottom: int = 35, end: int = 20) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total = sum(widths)
    tbl_w = table_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.top_margin = Inches(0.38)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.4)
    set_columns(section, 1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Inches(0.14)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    props = doc.core_properties
    props.title = "An End-to-End Lightweight Detection and Tracking Pipeline for UAV Imagery"
    props.subject = "VisDrone detection, UAV123 tracking, model compression, and TensorRT deployment"
    props.author = "Yukun Shi"
    props.keywords = "UAV vision, object detection, object tracking, small objects, TensorRT, knowledge distillation"
    props.comments = "All reported experimental values are generated from saved project results."
    return doc


def add_title_block(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("An End-to-End Lightweight Detection and Tracking Pipeline for UAV Imagery")
    set_run_font(run, 24)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.first_line_indent = None
    author.paragraph_format.space_after = Pt(1)
    set_run_font(author.add_run("Yukun Shi"), 11)

    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.first_line_indent = None
    affiliation.paragraph_format.space_after = Pt(4)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(4)
    set_run_font(paragraph.add_run(label), 9, bold=True, italic=True)
    set_run_font(paragraph.add_run(text), 9)


def add_body_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    set_columns(section, 2)


def add_heading_1(doc: Document, number: str, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    prefix = f"{number}.  " if number else ""
    run = paragraph.add_run(f"{prefix}{title.upper()}")
    set_run_font(run, 10)
    run.font.small_caps = True


def add_heading_2(doc: Document, letter: str, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(f"{letter}. {title}"), 10, italic=True)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.14)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    set_run_font(paragraph.add_run(text), 10)


def add_figure(doc: Document, image_path: Path, caption_number: int, caption: str, width: float = 3.34) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption_paragraph.paragraph_format.first_line_indent = None
    caption_paragraph.paragraph_format.space_after = Pt(4)
    set_run_font(caption_paragraph.add_run(f"Fig. {caption_number}. "), 8)
    set_run_font(caption_paragraph.add_run(caption), 8)


def add_table(
    doc: Document,
    table_number: str,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
) -> None:
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.first_line_indent = None
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after = Pt(1)
    title_p.paragraph_format.keep_with_next = True
    set_run_font(title_p.add_run(f"TABLE {table_number}"), 8)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = None
    subtitle.paragraph_format.space_after = Pt(2)
    subtitle.paragraph_format.keep_with_next = True
    set_run_font(subtitle.add_run(title.upper()), 8)

    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.first_line_indent = None
        set_run_font(cell.paragraphs[0].add_run(header), 7, bold=True)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E5E7EB")
        cell._tc.get_or_add_tcPr().append(shading)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_after = Pt(0)
            set_run_font(paragraph.add_run(value), 7)
    set_table_width(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.first_line_indent = None
    after.paragraph_format.space_after = Pt(2)


def add_references(doc: Document, references: list[str]) -> None:
    for index, reference in enumerate(references, start=1):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        set_run_font(paragraph.add_run(f"[{index}] {reference}"), 8)


def build_document() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    train_eda = load_json("outputs/eda/visdrone_train/summary.json")
    val_eda = load_json("outputs/eda/visdrone_val/summary.json")
    uav_eda = load_json("outputs/eda/uav123/summary.json")
    baseline = load_json("outputs/evaluation/yolov8s_visdrone_baseline_e10/summary.json")
    teacher = load_json("outputs/evaluation/yolov8s_visdrone_mildaug_e100/summary.json")
    nano = load_json("outputs/evaluation/yolov8n_visdrone_aug_e50/summary.json")
    slim = load_json("outputs/evaluation/yolov8s_slim04375_visdrone_e100/summary.json")
    distilled = load_json("outputs/evaluation/yolov8s_slim04375_distilled_e20/summary.json")
    mobilenet = load_json("outputs/evaluation/mobilenet_fpn_aerial_e10_complete/summary.json")
    mobilenet_balanced = load_json("outputs/training/mobilenet_fpn_visdrone_aerial_balanced_e12/summary.json")
    kcf = load_json("outputs/tracking/uav123_opencv_kcf/summary.json")["trackers"][0]
    csrt = load_json("outputs/tracking/uav123_opencv_csrt/summary.json")["trackers"][0]
    dasiam = load_json("outputs/tracking/uav123_opencv_dasiamrpn/summary.json")["trackers"][0]
    mot = load_json("outputs/tracking/yolov8s_slim04375_tensorrt_bytetrack_uav123_group1/summary.json")
    small_rows = load_csv("docs/assets/yolo_small_object_ablation/summary.csv")
    calibration = load_csv("docs/assets/int8_scene_calibration/calibration_results.csv")
    sensitivity = load_csv("docs/assets/int8_scene_calibration/sensitivity_results.csv")
    runtime = load_csv("docs/assets/runtime_memory_benchmark/summary.csv")

    pipeline_figure = build_pipeline_figure()
    eda_figure = build_eda_figure(train_eda, val_eda)
    small_figure = build_small_object_figure(small_rows)
    deployment_figure = build_deployment_figure(calibration, runtime)
    qualitative_figure = build_qualitative_figure()
    tracker_figure = ASSET_DIR / "tracker_comparison.png"
    shutil.copy2(ROOT / "outputs/tracking/uav123_opencv_tracker_comparison/tracker_comparison.png", tracker_figure)
    distillation_figure = ASSET_DIR / "distillation_comparison.png"
    shutil.copy2(ROOT / "docs/assets/yolo_distillation/distillation_comparison.png", distillation_figure)

    doc = configure_document()
    add_title_block(doc)
    add_labeled_paragraph(
        doc,
        "Abstract-",
        "This paper presents an end-to-end aerial vision pipeline covering dataset analysis, object detection, "
        "single- and multi-object tracking, model compression, small-object ablation, knowledge distillation, "
        "and GPU deployment. Experiments use the VisDrone detection split and all 123 UAV123 tracking sequences. "
        "A full YOLOv8s detector reached 0.2933 mAP50-95, while a width-0.4375 student reduced checkpoint size "
        "by 22.4% and retained 0.2696 mAP50-95. SAHI raised slim-model small-object recall from 0.4868 to 0.6456 "
        "but reduced throughput from 96.2 to 10.0 FPS. Output-distribution distillation produced a small "
        "heavy-occlusion recall gain in the auxiliary group analysis but did not improve small-object recall. On UAV123, DaSiamRPN "
        "achieved 0.5564 success AUC at 39.9 FPS. The final TensorRT FP16 detector provided the best deployment "
        "trade-off: 0.2700 mAP50-95, 91.7 end-to-end FPS, and 1175.9 MiB peak process memory. INT8 was slower and "
        "less accurate on the evaluated GPU. These results show that aerial accuracy is governed primarily by "
        "target scale and occlusion, while effective deployment requires joint optimization of architecture, "
        "inference strategy, numerical precision, and tracking overhead.",
    )
    add_labeled_paragraph(
        doc,
        "Index Terms-",
        "unmanned aerial vehicles, object detection, object tracking, small objects, model compression, "
        "knowledge distillation, TensorRT.",
    )
    add_body_section(doc)

    add_heading_1(doc, "I", "Introduction")
    add_body(
        doc,
        "UAV imagery combines wide fields of view with small targets, dense traffic, camera motion, and frequent "
        "occlusion. The VisDrone benchmark formalized these difficulties for detection and tracking [1], while "
        "UAV123 provides 123 low-altitude single-object tracking sequences [2]. A practical airborne system must "
        "therefore balance localization accuracy, temporal stability, model size, latency, and memory rather than "
        "optimize a single benchmark metric.",
    )
    add_body(
        doc,
        "The project developed a reproducible pipeline from raw annotation analysis to a deployable "
        "detection-and-tracking application. YOLOv8 [3] was selected as the principal one-stage detector. "
        "Faster R-CNN [4] with an FPN [5] and a MobileNetV3 backbone [6] was also evaluated as an alternative "
        "lightweight direction. The study then tested resizing and color augmentation, conservative width "
        "reduction, focal loss [7], sliced inference [8], single-object Siamese tracking [9], online data "
        "association [10], [11], knowledge distillation [12], ONNX interchange [13], and TensorRT deployment [14].",
    )
    add_body(
        doc,
        "The contributions are fourfold. First, the paper reports a complete dataset-to-deployment implementation "
        "using saved experimental artifacts rather than estimated values. Second, it isolates small-object and "
        "occlusion behavior with group-specific recall. Third, it compares classical and Siamese single-object "
        "trackers and integrates ByteTrack into the detector. Fourth, it documents negative results: aggressive "
        "slimming, untuned focal loss, image-level resampling, and the tested INT8 route did not improve the final "
        "accuracy-efficiency operating point.",
    )
    add_figure(doc, pipeline_figure, 1, "Implemented processing and deployment pipeline. The teacher model is used only during training.")

    add_heading_1(doc, "II", "Related Work")
    add_heading_2(doc, "A", "Aerial Detection and Small Objects")
    add_body(
        doc,
        "VisDrone includes urban and suburban scenes with object categories, bounding boxes, truncation, and "
        "occlusion annotations [1]. Multi-scale representation is commonly addressed by FPN [5]. Focal loss "
        "reduces the contribution of easy examples under class imbalance [7], while SAHI enlarges the effective "
        "pixel footprint of small targets by applying a detector to overlapping slices [8]. COCO-style AP over "
        "multiple IoU thresholds is used as a localization-sensitive summary metric [15].",
    )
    add_heading_2(doc, "B", "Tracking and Deployment")
    add_body(
        doc,
        "DaSiamRPN improves Siamese tracking robustness by explicitly learning distractor-aware features [9]. "
        "KCF and CSRT provide complementary classical correlation-filter baselines for speed-robustness comparison [16], [17]. "
        "For multi-object tracking, SORT demonstrated that detector quality strongly affects association quality "
        "[10], and ByteTrack later recovered low-score detections through a two-stage association strategy [11]. "
        "For deployment, ONNX provides a portable graph representation [13], whereas TensorRT performs graph "
        "optimization, layer fusion, kernel selection, and reduced-precision execution on NVIDIA GPUs [14].",
    )

    add_heading_1(doc, "III", "Data and Experimental Protocol")
    add_heading_2(doc, "A", "Datasets")
    add_body(
        doc,
        f"The VisDrone training split contains {train_eda['image_count']:,} images and "
        f"{train_eda['object_count']:,} annotated objects; the validation split contains "
        f"{val_eda['image_count']:,} images and {val_eda['object_count']:,} objects. Objects with bounding-box area below 32 x 32 pixels "
        f"account for {train_eda['small_object_ratio_lt_32x32']:.1%} of training annotations and "
        f"{val_eda['small_object_ratio_lt_32x32']:.1%} of validation annotations. UAV123 contains "
        f"{uav_eda['sequence_count']} sequences and {uav_eda['total_annotated_frames']:,} annotated frames.",
    )
    add_table(
        doc,
        "I",
        "Dataset Summary",
        ["Dataset", "Images / Sequences", "Objects / Frames"],
        [
            ["VisDrone train", f"{train_eda['image_count']:,} images", f"{train_eda['object_count']:,} objects"],
            ["VisDrone validation", f"{val_eda['image_count']:,} images", f"{val_eda['object_count']:,} objects"],
            ["UAV123", f"{uav_eda['sequence_count']} sequences", f"{uav_eda['total_annotated_frames']:,} frames"],
        ],
        [1800, 1500, 1500],
    )
    add_figure(doc, eda_figure, 2, "Measured class imbalance and small-object prevalence in VisDrone.")

    add_heading_2(doc, "B", "Metrics and Hardware")
    add_body(
        doc,
        "Detection reports precision, recall, AP50, mAP50-95, and class-aware recall at IoU 0.50 for small "
        "objects and heavy occlusion. Single-object tracking reports success AUC, precision at 20 pixels, mean "
        "IoU, center error, and FPS. Deployment benchmarks use batch one after warm-up. All principal experiments "
        "ran in Python 3.10 with PyTorch and CUDA on an NVIDIA GeForce RTX 5080 Laptop GPU. TensorRT experiments "
        "used version 11.1. Results from distinct evaluator protocols are kept in separate comparisons.",
    )

    add_heading_1(doc, "IV", "Detection, Compression, and Tracking Results")
    add_heading_2(doc, "A", "Detector Development")
    add_body(
        doc,
        "The initial 10-epoch YOLOv8s baseline achieved 0.2551 mAP50-95. Extending the mild-augmentation recipe "
        "to the selected long-run checkpoint increased mAP50-95 to 0.2933 and small-object recall to 0.5201. "
        "A YOLOv8n model trained for 50 epochs was compact but reached only 0.2053 mAP50-95. An initial width-0.375 "
        "slim YOLOv8s also underperformed at 0.1687 mAP50-95. Conservative width-0.4375 scaling plus 100-epoch "
        "recovery produced an 8.62 M-parameter, 16.70 MiB checkpoint with 0.2696 mAP50-95.",
    )
    add_table(
        doc,
        "II",
        "Unified VisDrone Detector Results (AP = mAP50-95)",
        ["Model", "P", "R", "AP50", "AP", "Small R"],
        [
            [
                "YOLOv8s e10",
                f"{baseline['training_metrics']['precision']:.3f}",
                f"{baseline['training_metrics']['recall']:.3f}",
                f"{baseline['training_metrics']['map50']:.3f}",
                f"{baseline['training_metrics']['map50_95']:.3f}",
                f"{baseline['groups']['small_lt_32x32']['recall_at_iou']:.3f}",
            ],
            [
                "YOLOv8n e50",
                f"{nano['training_metrics']['precision']:.3f}",
                f"{nano['training_metrics']['recall']:.3f}",
                f"{nano['training_metrics']['map50']:.3f}",
                f"{nano['training_metrics']['map50_95']:.3f}",
                f"{nano['groups']['small_lt_32x32']['recall_at_iou']:.3f}",
            ],
            [
                "YOLOv8s full",
                f"{teacher['training_metrics']['precision']:.3f}",
                f"{teacher['training_metrics']['recall']:.3f}",
                f"{teacher['training_metrics']['map50']:.3f}",
                f"{teacher['training_metrics']['map50_95']:.3f}",
                f"{teacher['groups']['small_lt_32x32']['recall_at_iou']:.3f}",
            ],
            [
                "Slim 0.4375",
                f"{slim['training_metrics']['precision']:.3f}",
                f"{slim['training_metrics']['recall']:.3f}",
                f"{slim['training_metrics']['map50']:.3f}",
                f"{slim['training_metrics']['map50_95']:.3f}",
                f"{slim['groups']['small_lt_32x32']['recall_at_iou']:.3f}",
            ],
            [
                "Distilled slim",
                f"{distilled['training_metrics']['precision']:.3f}",
                f"{distilled['training_metrics']['recall']:.3f}",
                f"{distilled['training_metrics']['map50']:.3f}",
                f"{distilled['training_metrics']['map50_95']:.3f}",
                f"{distilled['groups']['small_lt_32x32']['recall_at_iou']:.3f}",
            ],
        ],
        [1450, 650, 650, 650, 700, 700],
    )
    add_body(
        doc,
        "The MobileNet direction did not replace YOLO. The aerial MobileNetV3-FPN baseline obtained 0.1033 "
        "mAP50-95, 0.2205 small-object recall, and 78.1 FPS. Class-balanced focal loss, small anchors, and "
        "small-box weighting raised small-object recall to 0.2426 and heavy-occlusion recall from 0.1909 to "
        "0.2239, but reduced mAP50-95 to 0.0996 and throughput to 73.4 FPS. The result is a recall-oriented "
        "ablation, not a superior deployment model.",
    )

    add_heading_2(doc, "B", "Single- and Multi-Object Tracking")
    add_body(
        doc,
        "All image-based single-object trackers were evaluated on 123 UAV123 sequences. DaSiamRPN achieved the "
        "highest success AUC and precision, while KCF was the fastest but least accurate. CSRT occupied the middle "
        "operating point. The results support a learned Siamese tracker when robustness is more important than "
        "maximum throughput.",
    )
    add_table(
        doc,
        "III",
        "UAV123 Single-Object Tracking",
        ["Tracker", "AUC", "P@20", "Mean IoU", "FPS"],
        [
            ["DaSiamRPN", f"{dasiam['success_auc']:.3f}", f"{dasiam['precision_20']:.3f}", f"{dasiam['mean_iou']:.3f}", f"{dasiam['fps']:.1f}"],
            ["CSRT", f"{csrt['success_auc']:.3f}", f"{csrt['precision_20']:.3f}", f"{csrt['mean_iou']:.3f}", f"{csrt['fps']:.1f}"],
            ["KCF", f"{kcf['success_auc']:.3f}", f"{kcf['precision_20']:.3f}", f"{kcf['mean_iou']:.3f}", f"{kcf['fps']:.1f}"],
        ],
        [1450, 820, 820, 900, 810],
    )
    add_figure(doc, tracker_figure, 3, "Full UAV123 comparison of the three image-based single-object trackers.")
    add_body(
        doc,
        f"The final detector-to-ByteTrack pipeline processed {mot['processed_frames']} frames from UAV123 group1, "
        f"produced {mot['total_rows']:,} track rows and {mot['unique_track_ids']} unique IDs, and ran at "
        f"{mot['fps']:.2f} FPS with TensorRT FP16. Because this sequence lacks complete multi-object identity "
        "ground truth, track counts and throughput are reported without inventing MOTA, IDF1, or HOTA values.",
    )
    add_figure(doc, qualitative_figure, 4, "Real validation predictions and integrated tracking output produced by the implemented pipeline.")

    add_heading_1(doc, "V", "Small-Object and Distillation Ablations")
    add_heading_2(doc, "A", "Small-Object Interventions")
    standard = next(row for row in small_rows if row["key"] == "standard")
    multiscale = next(row for row in small_rows if row["key"] == "multiscale_768_960_1280")
    sahi = next(row for row in small_rows if row["key"] == "sahi_640_overlap020")
    focal = next(row for row in small_rows if row["key"] == "focal_gamma2_e6_eval")
    resample = next(row for row in small_rows if row["key"] == "small_resample_strength2_e6_eval")
    add_body(
        doc,
        f"On the slim detector, multi-scale inference increased small-object recall from "
        f"{float(standard['small_recall']):.4f} to {float(multiscale['small_recall']):.4f} at "
        f"{float(multiscale['fps']):.2f} FPS. SAHI reached {float(sahi['small_recall']):.4f} small-object recall "
        f"and {float(sahi['map50']):.4f} mAP50, but throughput fell to {float(sahi['fps']):.2f} FPS. Focal loss "
        f"raised recall to {float(focal['small_recall']):.4f} while slightly reducing mAP50-95. Image-level "
        f"resampling produced {float(resample['small_recall']):.4f} recall, essentially no benefit. Thus SAHI is "
        "an offline accuracy mode, whereas standard inference remains the real-time mode.",
    )
    add_figure(doc, small_figure, 5, "Measured small-object recall and throughput for selected slim-detector ablations.")

    add_heading_2(doc, "B", "Teacher-Student Distillation")
    add_body(
        doc,
        "A full YOLOv8s teacher supervised aligned class logits and 16-bin bounding-box distributions from the "
        "slim student. The 20-epoch run preserved the student architecture. Recall increased from 0.4585 to "
        "0.4640 and heavy-occlusion recall from 0.3611 to 0.3702 in the saved auxiliary group-recall summaries. "
        "However, mAP50-95 changed only from 0.2696 to 0.2699, and small-object recall "
        "changed from 0.4955 to 0.4953. Output-only distillation therefore helped occlusion modestly but did not "
        "transfer enough high-resolution feature information to improve tiny targets.",
    )
    add_figure(doc, distillation_figure, 6, "Teacher, slim baseline, and distilled slim comparison from saved evaluation summaries.")

    add_heading_1(doc, "VI", "Deployment Results")
    fp16_cal = next(row for row in calibration if row["method"] == "FP16")
    int8_bright = next(row for row in calibration if row["method"] == "INT8 bright")
    head_fp16 = next(row for row in sensitivity if row["method"] == "Head FP16")
    rt = {row["name"]: row for row in runtime}
    add_body(
        doc,
        "PyTorch checkpoints were exported to ONNX and checked against framework predictions before TensorRT "
        "engine construction. In the full 548-image engine validation, FP16 achieved 0.2700 mAP50-95 and 477.2 "
        "raw engine FPS. Bright-scene INT8 calibration was the best of three calibration sets but reached only "
        "0.2467 mAP50-95 and 333.9 FPS. Restoring the detection head to FP16 recovered mAP50-95 to 0.2617, "
        "confirming that classification, DFL, and decode layers were quantization-sensitive, but the hybrid still "
        "underperformed full FP16.",
    )
    add_table(
        doc,
        "IV",
        "TensorRT Accuracy and Engine Throughput (AP = mAP50-95)",
        ["Engine", "P", "R", "AP50", "AP", "Raw FPS"],
        [
            ["FP16", f"{float(fp16_cal['precision']):.3f}", f"{float(fp16_cal['recall']):.3f}", f"{float(fp16_cal['map50']):.3f}", f"{float(fp16_cal['map50_95']):.3f}", f"{float(fp16_cal['fps']):.1f}"],
            ["INT8 bright", f"{float(int8_bright['precision']):.3f}", f"{float(int8_bright['recall']):.3f}", f"{float(int8_bright['map50']):.3f}", f"{float(int8_bright['map50_95']):.3f}", f"{float(int8_bright['fps']):.1f}"],
            ["Head FP16", f"{float(head_fp16['precision']):.3f}", f"{float(head_fp16['recall']):.3f}", f"{float(head_fp16['map50']):.3f}", f"{float(head_fp16['map50_95']):.3f}", f"{float(head_fp16['fps']):.1f}"],
        ],
        [1450, 650, 650, 650, 700, 700],
    )
    add_body(
        doc,
        "An isolated end-to-end runtime benchmark, which includes the Ultralytics prediction path and "
        "post-processing, measured 91.7 FPS and 1175.9 MiB peak process RAM for TensorRT FP16. INT8 was both "
        "slower at 78.6 FPS and larger in memory at 1192.5 MiB. Adding ByteTrack reduced throughput to 63.8 FPS "
        "and increased peak RAM to 1215.0 MiB. The difference from raw engine FPS is expected because the "
        "end-to-end protocol includes preprocessing, result construction, NMS, and tracking.",
    )
    add_figure(doc, deployment_figure, 7, "TensorRT validation accuracy and end-to-end throughput. The two panels use their stated protocols.")

    add_heading_1(doc, "VII", "Discussion and Limitations")
    add_body(
        doc,
        "The experiments identify a stable hierarchy of constraints. First, spatial resolution dominates "
        "small-object recall: multi-scale and sliced inference gave the largest gains, but every gain required "
        "multiple forward passes. Second, moderate width reduction preserved much more accuracy than aggressive "
        "slimming, showing that recovery training cannot compensate for an undersized feature hierarchy. Third, "
        "tracking quality remained detector-dependent; ByteTrack made the pipeline usable, but its temporal "
        "accuracy cannot be claimed without identity-complete UAV ground truth. Fourth, reduced precision was "
        "hardware- and graph-dependent. INT8 did not automatically improve latency and was especially sensitive "
        "in the detection head.",
    )
    add_body(
        doc,
        "Several limitations bound the conclusions. All compute measurements were made on one laptop GPU rather "
        "than a Jetson-class target, and WDDM prevented reliable per-process VRAM attribution. VisDrone validation "
        "metrics are produced by project evaluators and are not official challenge-server test scores. Some "
        "ablation tables use different score-retention rules and are compared only internally. The current "
        "distillation loss operates on output distributions rather than ground-truth-masked P3 features. Finally, "
        "the multi-object demonstration measures engineering completeness and throughput, not MOTA or IDF1.",
    )

    add_heading_1(doc, "VIII", "Conclusion")
    add_body(
        doc,
        "The completed system demonstrates a reproducible UAV perception workflow from data understanding to "
        "deployment. Full YOLOv8s provided the best detector accuracy, while the width-0.4375 student offered a "
        "useful storage-accuracy compromise. DaSiamRPN was the strongest evaluated single-object tracker, and "
        "ByteTrack enabled an integrated multi-object pipeline. SAHI substantially improved small-target recall "
        "but was not real-time. Knowledge distillation partially recovered occlusion recall without changing the "
        "student architecture. TensorRT FP16 was the final deployment choice because it outperformed the tested "
        "INT8 engines in accuracy, speed, and memory. Future work should evaluate P3 feature distillation, "
        "adaptive slicing, identity-labeled aerial MOT sequences, and sustained thermal performance on the actual "
        "embedded device.",
    )

    add_heading_1(doc, "", "References")
    references = [
        'P. Zhu, L. Wen, X. Bian, H. Ling, and Q. Hu, "Vision Meets Drones: A Challenge," arXiv:1804.07437, 2018.',
        'M. Mueller, N. Smith, and B. Ghanem, "A Benchmark and Simulator for UAV Tracking," in Proc. ECCV, 2016, pp. 445-461, doi: 10.1007/978-3-319-46448-0_27.',
        'G. Jocher, A. Chaurasia, and J. Qiu, "Ultralytics YOLOv8," version 8.0.0, software, 2023. [Online]. Available: https://github.com/ultralytics/ultralytics',
        'S. Ren, K. He, R. Girshick, and J. Sun, "Faster R-CNN: Towards Real-Time Object Detection with Region Proposal Networks," in Adv. Neural Inf. Process. Syst., vol. 28, 2015.',
        'T.-Y. Lin, P. Dollar, R. Girshick, K. He, B. Hariharan, and S. Belongie, "Feature Pyramid Networks for Object Detection," in Proc. IEEE CVPR, 2017, pp. 2117-2125.',
        'A. Howard et al., "Searching for MobileNetV3," in Proc. IEEE/CVF ICCV, 2019, pp. 1314-1324.',
        'T.-Y. Lin, P. Goyal, R. Girshick, K. He, and P. Dollar, "Focal Loss for Dense Object Detection," in Proc. IEEE ICCV, 2017, pp. 2980-2988.',
        'F. C. Akyon, S. O. Altinuc, and A. Temizel, "Slicing Aided Hyper Inference and Fine-Tuning for Small Object Detection," in Proc. IEEE ICIP, 2022, pp. 966-970, doi: 10.1109/ICIP46576.2022.9897990.',
        'Z. Zhu, Q. Wang, B. Li, W. Wu, J. Yan, and W. Hu, "Distractor-Aware Siamese Networks for Visual Object Tracking," in Proc. ECCV, 2018, pp. 101-117.',
        'A. Bewley, Z. Ge, L. Ott, F. Ramos, and B. Upcroft, "Simple Online and Realtime Tracking," in Proc. IEEE ICIP, 2016, pp. 3464-3468, doi: 10.1109/ICIP.2016.7533003.',
        'Y. Zhang et al., "ByteTrack: Multi-Object Tracking by Associating Every Detection Box," in Proc. ECCV, 2022, pp. 1-21, doi: 10.1007/978-3-031-20047-2_1.',
        'G. Hinton, O. Vinyals, and J. Dean, "Distilling the Knowledge in a Neural Network," arXiv:1503.02531, 2015.',
        'ONNX Working Group, "Open Neural Network Exchange Intermediate Representation Specification," version 1.23.0, 2026. [Online]. Available: https://onnx.ai/onnx/repo-docs/IR.html. Accessed: Jul. 28, 2026',
        'NVIDIA, "NVIDIA TensorRT Documentation," version 11.1.0, 2026. [Online]. Available: https://docs.nvidia.com/deeplearning/tensorrt/latest/. Accessed: Jul. 28, 2026',
        'T.-Y. Lin et al., "Microsoft COCO: Common Objects in Context," in Proc. ECCV, 2014, pp. 740-755.',
        'J. F. Henriques, R. Caseiro, P. Martins, and J. Batista, "High-Speed Tracking with Kernelized Correlation Filters," IEEE Trans. Pattern Anal. Mach. Intell., vol. 37, no. 3, pp. 583-596, 2015, doi: 10.1109/TPAMI.2014.2345390.',
        'A. Lukezic, T. Vojir, L. Cehovin Zajc, J. Matas, and M. Kristan, "Discriminative Correlation Filter with Channel and Spatial Reliability," in Proc. IEEE CVPR, 2017, pp. 6309-6318.',
    ]
    add_references(doc, references)

    balancing_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(balancing_section, 1)

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_document())









